from pdf2image import convert_from_path
import easyocr
from docx import Document
import numpy as np

def pdf_to_word_easyocr(pdf_path, output_docx):
    images = convert_from_path(pdf_path, dpi=300, poppler_path=r"C:\Program Files\poppler\poppler-24.08.0\bin")
    reader = easyocr.Reader(['tr'])  # Türkçe desteği
    doc = Document()

    for i, image in enumerate(images):
        result = reader.readtext(np.array(image), detail=0, paragraph=True)
        text = "\n".join(result)
        doc.add_paragraph(f"--- Sayfa {i+1} ---")
        doc.add_paragraph(text)

    doc.save(output_docx)
    print(f"Dönüştürüldü: {output_docx}")

# Kullanım
pdf_to_word_easyocr("girdi.pdf", "cikti_easyocr.docx")